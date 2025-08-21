from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
import base64
import os
from typing import List, Optional
import asyncio
from google import genai
from google.genai import types
from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.config import Settings
import uuid
import io
import PyPDF2
from docx import Document
import hashlib
from dotenv import load_dotenv

load_dotenv("../../.env")

# Create uploads directory if it doesn't exist
UPLOADS_DIR = "./uploads"
os.makedirs(UPLOADS_DIR, exist_ok=True)

app = FastAPI(title="Invoice Document Processing API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],         # In production, specify frontend domain(s) here for best security.
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Gemini client
try:
    client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))
except Exception as e:
    print(f"Warning: Gemini client initialization failed: {e}")
    client = None

# Initialize SentenceTransformer model for embeddings
try:
    embedder = SentenceTransformer('all-MiniLM-L6-v2')
except Exception as e:
    print(f"Warning: SentenceTransformer initialization failed: {e}")
    embedder = None

# Initialize ChromaDB client (persistent storage)
try:
    chroma_client = chromadb.PersistentClient(path="./chroma_db")
    collection_name = "invoice_documents"
    
    # Create or get collection with custom embedding function
    try:
        collection = chroma_client.get_collection(collection_name)
    except:
        collection = chroma_client.create_collection(
            name=collection_name,
            metadata={"hnsw:space": "cosine"}
        )
except Exception as e:
    print(f"Warning: ChromaDB initialization failed: {e}")
    collection = None

class DocumentProcessor:
    def __init__(self):
        self.gemini_client = client
        self.embedder = embedder
        
    async def extract_text_from_pdf_gemini(self, file_bytes: bytes, filename: str) -> dict:
        """Extract information from PDF using Gemini LLM"""
        try:
            model = "gemini-2.0-flash"
            contents = [
                types.Content(
                    role="user",
                    parts=[
                        types.Part.from_bytes(
                            mime_type="application/pdf", 
                            data=file_bytes
                        ),
                        types.Part.from_text(
                            text="""Extract all relevant information from this invoice/document. 
                            Please provide a structured summary including:
                            - Document type
                            - Invoice number (if applicable)
                            - Date
                            - Vendor/Company information
                            - Customer information
                            - Items/services listed
                            - Amounts and totals
                            - Any other relevant details
                            
                            Format the response in a clear, structured way."""
                        ),
                    ],
                )
            ]
            
            result_text = ""
            generate_content_config = types.GenerateContentConfig()
            
            for chunk in self.gemini_client.models.generate_content_stream(
                model=model, 
                contents=contents, 
                config=generate_content_config
            ):
                result_text += chunk.text
            
            return {
                "success": True,
                "extracted_text": result_text,
                "method": "gemini_llm"
            }
        except Exception as e:
            # Fallback to basic PDF extraction
            return await self.extract_text_from_pdf_fallback(file_bytes)
    
    async def extract_text_from_pdf_fallback(self, file_bytes: bytes) -> dict:
        """Fallback PDF text extraction using PyPDF2"""
        try:
            pdf_file = io.BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return {
                "success": True,
                "extracted_text": text,
                "method": "pypdf2_fallback"
            }
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "method": "fallback_failed"
            }
    
    async def extract_text_from_docx(self, file_bytes: bytes) -> dict:
        """Extract text from Word document"""
        try:
            doc_file = io.BytesIO(file_bytes)
            doc = Document(doc_file)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return {
                "success": True,
                "extracted_text": text,
                "method": "python_docx"
            }
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "method": "docx_failed"
            }
    
    def generate_embedding(self, text: str) -> List[float]:
        """Generate embedding for text using sentence-transformers"""
        try:
            return self.embedder.encode(text).tolist()
        except Exception as e:
            # Return zero vector as fallback
            return [0.0] * 384  # all-MiniLM-L6-v2 has 384 dimensions

processor = DocumentProcessor()

@app.get("/")
async def root():
    return {"message": "Invoice Document Processing API", "status": "running"}

@app.get("/health")
async def health_check():
    return {
        "gemini_available": client is not None,
        "embedder_available": embedder is not None,
        "chromadb_available": collection is not None
    }

@app.post("/upload-document/")
async def upload_document(file: UploadFile = File(...)):
    """Upload and process invoice/document"""
    try:
        if not collection:
            raise HTTPException(status_code=500, detail="ChromaDB not available")
        
        # Read file content
        file_bytes = await file.read()
        file_hash = hashlib.md5(file_bytes).hexdigest()
        doc_id = f"{file.filename}_{file_hash[:8]}"
        
        # Save the original file
        file_path = os.path.join(UPLOADS_DIR, f"{doc_id}_{file.filename}")
        with open(file_path, "wb") as f:
            f.write(file_bytes)
        
        # Check if document already exists
        try:
            existing = collection.get(ids=[doc_id])
            if existing['ids']:
                return {
                    "message": "Document already exists",
                    "doc_id": doc_id,
                    "status": "duplicate"
                }
        except:
            pass
        
        # Extract text based on file type
        extraction_result = None
        
        if file.content_type == "application/pdf":
            extraction_result = await processor.extract_text_from_pdf_gemini(
                file_bytes, file.filename
            )
        elif file.content_type in [
            "application/msword", 
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ]:
            extraction_result = await processor.extract_text_from_docx(file_bytes)
        else:
            raise HTTPException(
                status_code=415, 
                detail=f"Unsupported file type: {file.content_type}"
            )
        
        if not extraction_result["success"]:
            raise HTTPException(
                status_code=500, 
                detail=f"Text extraction failed: {extraction_result.get('error', 'Unknown error')}"
            )
        
        extracted_text = extraction_result["extracted_text"]
        
        # Generate embedding
        embedding = processor.generate_embedding(extracted_text)
        
        # Store in vector database
        collection.add(
            documents=[extracted_text],
            metadatas=[{
                "filename": file.filename,
                "content_type": file.content_type,
                "extraction_method": extraction_result["method"],
                "file_size": len(file_bytes),
                "file_path": file_path
            }],
            ids=[doc_id],
            embeddings=[embedding]
        )
        
        return {
            "message": "Document uploaded and processed successfully",
            "doc_id": doc_id,
            "extraction_method": extraction_result["method"],
            "text_length": len(extracted_text),
            "status": "success"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    


# @app.post("/query-documents/")
# async def query_documents(
#     query: str, 
#     top_k: int = 3,
#     use_llm_for_answer: bool = True
# ):
#     """Query documents and get relevant answers"""
#     try:
#         if not collection:
#             raise HTTPException(status_code=500, detail="ChromaDB not available")
        
#         # Generate embedding for query
#         query_embedding = processor.generate_embedding(query)
        
#         # Query vector database
#         results = collection.query(
#             query_embeddings=[query_embedding],
#             n_results=top_k,
#             include=['documents']  # Only include documents to avoid metadata issues
#         )
        
#         if not results['documents'] or not results['documents'][0]:
#             return {
#                 "query": query,
#                 "answer": "No relevant documents found.",
#                 "sources": [],
#                 "method": "no_results"
#             }
        
#         # Extract documents
#         documents = results['documents'][0]
        
#         # If Gemini is available and requested, use it to generate a comprehensive answer
#         if use_llm_for_answer and client:
#             try:
#                 context = "\n\n".join(documents[:2])  # Use top 2 most relevant docs
                
#                 model = "gemini-2.0-flash"
#                 contents = [
#                     types.Content(
#                         role="user",
#                         parts=[
#                             types.Part.from_text(
#                                 text=f"""Based on the following document excerpts, please answer this question: "{query}"
                                
#                                 Document excerpts:
#                                 {context}
                                
#                                 Please provide a clear, accurate answer based only on the information provided in the documents. 
#                                 If the information is not sufficient to answer the question, please say so."""
#                             ),
#                         ],
#                     )
#                 ]
                
#                 llm_answer = ""
#                 generate_content_config = types.GenerateContentConfig()
                
#                 for chunk in client.models.generate_content_stream(
#                     model=model, 
#                     contents=contents, 
#                     config=generate_content_config
#                 ):
#                     llm_answer += chunk.text
                
#                 return {
#                     "query": query,
#                     "answer": llm_answer,
#                     "sources": [
#                         {
#                             "document_index": i + 1,
#                             "content_preview": doc[:300] + "..." if len(doc) > 300 else doc
#                         }
#                         for i, doc in enumerate(documents)
#                     ],
#                     "method": "llm_generated"
#                 }
                
#             except Exception as e:
#                 print(f"LLM generation failed: {e}")
#                 pass
        
#         # Fallback: return relevant document excerpts
#         return {
#             "query": query,
#             "answer": "Here are the most relevant document excerpts:",
#             "sources": [
#                 {
#                     "document_index": i + 1,
#                     "content": doc[:500] + "..." if len(doc) > 500 else doc
#                 }
#                 for i, doc in enumerate(documents)
#             ],
#             "method": "document_retrieval"
#         }
        
#     except HTTPException:
#         raise
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))
from pydantic import BaseModel
from typing import Optional

class QueryRequest(BaseModel):
    query: str
    filename: Optional[str] = None
    top_k: Optional[int] = 3
    use_llm_for_answer: Optional[bool] = True

# Replace your current endpoint with this:
@app.post("/query-documents/")
async def query_documents(request: QueryRequest):
    """Query documents and get relevant answers"""
    try:
        if not collection:
            raise HTTPException(status_code=500, detail="ChromaDB not available")
        
        # Generate embedding for query
        query_embedding = processor.generate_embedding(request.query)
        
        # Query vector database

        results = None
        if request.filename:
            results = collection.query(
                query_embeddings=[query_embedding],
                n_results=request.top_k,
                include=['documents'],
                where={"filename": request.filename}
            )
        else:
            results = collection.query(
                query_embeddings=[query_embedding],
                n_results=request.top_k,
                include=['documents']
            )
        
        if not results['documents'] or not results['documents'][0]:
            return {
                "query": request.query,
                "answer": "No relevant documents found.",
                "sources": [],
                "method": "no_results"
            }
        
        # Extract documents
        documents = results['documents'][0]
        
        # If Gemini is available and requested, use it to generate a comprehensive answer
        if request.use_llm_for_answer and client:
            try:
                context = "\n\n".join(documents[:2])  # Use top 2 most relevant docs
                
                model = "gemini-2.0-flash"
                contents = [
                    types.Content(
                        role="user",
                        parts=[
                            types.Part.from_text(
                                text=f"""Based on the following document excerpts, please answer this question: "{request.query}"
                                
                                Document excerpts:
                                {context}
                                
                                Please provide a clear, accurate answer based only on the information provided in the documents. 
                                If the information is not sufficient to answer the question, please say so."""
                            ),
                        ],
                    )
                ]
                
                llm_answer = ""
                generate_content_config = types.GenerateContentConfig()
                
                for chunk in client.models.generate_content_stream(
                    model=model, 
                    contents=contents, 
                    config=generate_content_config
                ):
                    llm_answer += chunk.text
                
                return {
                    "query": request.query,
                    "answer": llm_answer,
                    "sources": [
                        {
                            "document_index": i + 1,
                            "content_preview": doc[:300] + "..." if len(doc) > 300 else doc
                        }
                        for i, doc in enumerate(documents)
                    ],
                    "method": "llm_generated"
                }
                
            except Exception as e:
                print(f"LLM generation failed: {e}")
                pass
        
        # Fallback: return relevant document excerpts
        return {
            "query": request.query,
            "answer": "Here are the most relevant document excerpts:",
            "sources": [
                {
                    "document_index": i + 1,
                    "content": doc[:500] + "..." if len(doc) > 500 else doc
                }
                for i, doc in enumerate(documents)
            ],
            "method": "document_retrieval"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/list-documents/")
async def list_documents():
    """List all uploaded documents"""
    try:
        if not collection:
            raise HTTPException(status_code=500, detail="ChromaDB not available")
        
        # Get all documents
        all_docs = collection.get(include=['metadatas'])
        
        documents = [
            {
                "doc_id": doc_id,
                "filename": metadata.get('filename', 'Unknown'),
                "content_type": metadata.get('content_type', 'Unknown'),
                "extraction_method": metadata.get('extraction_method', 'Unknown'),
                "file_size": metadata.get('file_size', 0)
            }
            for doc_id, metadata in zip(all_docs['ids'], all_docs['metadatas'])
        ]
        
        return {
            "total_documents": len(documents),
            "documents": documents
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/view-document/{doc_id}")
async def view_document(doc_id: str):
    """Get the original document file for viewing"""
    try:
        if not collection:
            raise HTTPException(status_code=500, detail="ChromaDB not available")
        
        # Get document metadata
        result = collection.get(ids=[doc_id], include=["metadatas"])
        
        if not result['ids']:
            raise HTTPException(status_code=404, detail="Document not found")
        
        metadata = result['metadatas'][0]
        file_path = metadata.get('file_path')
        
        if not file_path or not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="Original file not found")
        
        # Create response with inline disposition for PDFs
        response = FileResponse(
            path=file_path,
            media_type=metadata['content_type'],
            filename=metadata['filename']
        )
        
        # For PDFs, set Content-Disposition to inline to display in browser
        if metadata['content_type'] == 'application/pdf':
            response.headers["Content-Disposition"] = f'inline; filename="{metadata["filename"]}"'
        
        return response
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/delete-document/{doc_id}")
async def delete_document(doc_id: str):
    """Delete a specific document"""
    try:
        if not collection:
            raise HTTPException(status_code=500, detail="ChromaDB not available")
        
        # Get document metadata to find file path
        try:
            result = collection.get(ids=[doc_id], include=["metadatas"])
            if result['ids'] and result['metadatas']:
                metadata = result['metadatas'][0]
                file_path = metadata.get('file_path')
                
                # Delete the physical file if it exists
                if file_path and os.path.exists(file_path):
                    os.remove(file_path)
        except:
            pass  # Continue with deletion even if file removal fails
        
        # Delete from vector database
        collection.delete(ids=[doc_id])
        
        return {
            "message": f"Document {doc_id} deleted successfully",
            "status": "success"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
